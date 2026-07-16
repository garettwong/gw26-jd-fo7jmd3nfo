from __future__ import annotations

from copy import deepcopy
from pathlib import Path
import shutil

from docx import Document
from docx.shared import RGBColor


ROOT = Path(r"D:/Garett Super Jobs 2026/Calvin/REAL ERB/Check schedule only")
SOURCE = ROOT / "HK239HG(CW10)_R3_CONFIRMED FINAL.docx"
OUTPUT = ROOT / "checked" / "HK280HSSS_R3_20260716_V14_CHECKED.docx"


def set_paragraph_text(paragraph, text: str, *, red: bool = False) -> None:
    source_rpr = deepcopy(paragraph.runs[0]._r.rPr) if paragraph.runs else None
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)
    run = paragraph.add_run(text)
    if source_rpr is not None:
        run._r.insert(0, source_rpr)
    if red:
        run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
        run.bold = True


def set_cell_text(cell, text: str, *, red: bool = False) -> None:
    set_paragraph_text(cell.paragraphs[0], text, red=red)
    for extra in list(cell.paragraphs[1:]):
        cell._tc.remove(extra._p)


def replace_paragraph(document: Document, starts_with: str, replacement: str) -> None:
    for paragraph in document.paragraphs:
        if paragraph.text.strip().startswith(starts_with):
            set_paragraph_text(paragraph, replacement)
            return
    raise ValueError(f"Missing paragraph starting with {starts_with!r}")


def create_reply() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(SOURCE, OUTPUT)
    document = Document(OUTPUT)

    replace_paragraph(document, "課程編號", "課程編號：HK280HS   班別：SS")
    replace_paragraph(document, "課程名稱", "課程名稱：生成式人工智能商務應用證書（兼讀制）")
    replace_paragraph(
        document,
        "上課地點",
        "上課地點：基督教勵行會，上水彩園邨彩湖樓2座地下129舖02室",
    )
    replace_paragraph(document, "上課日期", "上課日期：待Calvin安排（只可考慮2026年9月14日上午）")
    replace_paragraph(document, "課    節", "課    節：5節（全課共18小時）")

    schedule = document.tables[1]
    rows = [
        ("1", "9月14日", "一", "0900-1300", "", "可考慮（未確認）"),
        ("2", "—", "—", "0900-1230", "", "暫未能安排"),
        ("3", "—", "—", "0900-1230", "", "暫未能安排；持續評估-個人習作一"),
        ("4", "—", "—", "0900-1230", "", "暫未能安排；持續評估-個人習作二"),
        ("5", "—", "—", "0900-1230", "", "暫未能安排；期末筆試1100-1200"),
    ]
    for row_index, values in enumerate(rows, 1):
        for column_index, value in enumerate(values):
            set_cell_text(
                schedule.rows[row_index].cells[column_index],
                value,
                red=column_index == 5,
            )
    schedule._tbl.remove(schedule.rows[6]._tr)

    set_paragraph_text(
        document.paragraphs[11],
        "V14回覆：只可考慮9月14日上午，現階段未確認；該時段與HK281DS CW7未確認課節重疊，請確認HK280HS後才釋放原有保留。",
        red=True,
    )
    set_paragraph_text(
        document.paragraphs[12],
        "其餘不合適：9月14日下午（往彩雲晚課沒有可靠用膳時間）；9月15、18、22日（SEN）；9月16、17日下午（HK265HG）；9月21日上午（上水至四海需約64分鐘，趕不上14:00課堂），下午亦已有課。",
        red=True,
    )
    document.core_properties.comments = (
        "V14 enquiry reply prepared from the HK280HSSS_R3 screenshots and Calvin WhatsApp message dated 2026-07-16. "
        "Only 2026-09-14 AM is proposed and remains unconfirmed."
    )
    document.save(OUTPUT)


if __name__ == "__main__":
    create_reply()
    print(OUTPUT)
