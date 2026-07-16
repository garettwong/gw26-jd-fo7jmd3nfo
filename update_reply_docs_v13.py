from __future__ import annotations

from copy import deepcopy
from pathlib import Path
import shutil

from docx import Document
from docx.shared import RGBColor


CHECKED = Path(r"D:/Garett Super Jobs 2026/Calvin/REAL ERB/Check schedule only/checked")
CW8_SOURCE = CHECKED / "HK244HG(CW8)_R3_final_checked.docx"
CW8_OUTPUT = CHECKED / "HK244HG(CW8)_R3_final_checked_20260716_V13.docx"
FS_OUTPUT = CHECKED / "HK239HG(FS)8月班_0508_20260716_V13_CONFIRMED.docx"


def set_paragraph_text(paragraph, text: str, *, red: bool = False) -> None:
    source_rpr = None
    if paragraph.runs:
        source_rpr = deepcopy(paragraph.runs[0]._r.rPr)
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)
    run = paragraph.add_run(text)
    if source_rpr is not None:
        run._r.insert(0, source_rpr)
    if red:
        run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
        run.bold = True


def set_cell_text(cell, text: str, *, red: bool = False) -> None:
    paragraph = cell.paragraphs[0]
    set_paragraph_text(paragraph, text, red=red)
    for extra in list(cell.paragraphs[1:]):
        cell._tc.remove(extra._p)


def replace_paragraph(document: Document, starts_with: str, replacement: str) -> None:
    for paragraph in document.paragraphs:
        if paragraph.text.strip().startswith(starts_with):
            set_paragraph_text(paragraph, replacement)
            return
    raise ValueError(f"Missing paragraph starting with {starts_with!r}")


def update_cw8_reply() -> None:
    shutil.copy2(CW8_SOURCE, CW8_OUTPUT)
    document = Document(CW8_OUTPUT)
    schedule = document.tables[1]
    set_cell_text(
        schedule.rows[3].cells[5],
        "Calvin代課（2026-07-16 WhatsApp確認）",
        red=True,
    )
    document.core_properties.comments = (
        "V13: HK244HG CW8 L3 on 2026-08-14 is taught by Calvin, who substitutes for Garett."
    )
    document.save(CW8_OUTPUT)


def create_fs_reply() -> None:
    shutil.copy2(CW8_SOURCE, FS_OUTPUT)
    document = Document(FS_OUTPUT)
    replace_paragraph(document, "課程編號", "課程編號：HK239HG   班別：FS")
    replace_paragraph(document, "課程名稱", "課程名稱︰人工智能知識及應用證書 (兼讀制)")
    replace_paragraph(
        document,
        "上課地點",
        "上課地點 : 九龍彌敦道208至212號四海大廈2樓全層（佐敦港鐵站D出口）",
    )
    replace_paragraph(document, "上課日期", "上課日期：2026年8月14日 至 2026年8月21日 (詳見附表)")
    replace_paragraph(document, "課堂課節", "課堂課節：6節 (全課共18小時)")

    schedule = document.tables[1]
    rows = [
        ("1", "8月14日", "五", "1000-1300", "3", "Garett"),
        ("2", "8月14日", "五", "1400-1700", "3", "Garett"),
        ("3", "8月19日", "三", "1000-1300", "3", "Garett"),
        ("4", "8月19日", "三", "1400-1700", "3", "Garett"),
        ("5", "8月21日", "五", "1000-1300", "3", "其他導師（持續評估／小組討論及專題報告）"),
        ("6", "8月21日", "五", "1400-1700", "3", "其他導師（期末筆試1530-1630）"),
    ]
    for row_index, values in enumerate(rows, 1):
        for column_index, value in enumerate(values):
            set_cell_text(
                schedule.rows[row_index].cells[column_index],
                value,
                red=column_index == 5,
            )
    for row in list(schedule.rows[7:]):
        schedule._tbl.remove(row._tr)

    confirmation = document.paragraphs[14]
    set_paragraph_text(
        confirmation,
        "V13確認：Garett教授第1至4節；第5至6節由其他導師教授。Calvin於2026年7月16日WhatsApp確認。",
        red=True,
    )
    document.core_properties.comments = (
        "Created from the HK239HG(FS)8月班_0508 screenshot and Calvin WhatsApp confirmation dated 2026-07-16."
    )
    document.save(FS_OUTPUT)


if __name__ == "__main__":
    update_cw8_reply()
    create_fs_reply()
    print(CW8_OUTPUT)
    print(FS_OUTPUT)
