from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.shared import RGBColor


ROOT = Path(r"D:\Garett Super Jobs 2026\Calvin\REAL ERB")
CONFIRMED = ROOT / "Check schedule only (Codex)" / "_____05 Confirmed Schedules"
FINAL = CONFIRMED / "00 FINAL FINAL - Calvin Confirmed"
SUPERSEDED = (
    ROOT
    / "Check schedule only (Codex)"
    / "04 Superseded - Do Not Send"
    / "20260811 V20k teacher-allocation error"
)
HK239_CONFIRMATION = (
    "最新確認：Garett 教授第 1、2、3、4 節；第 5 節由 Judy 教授，"
    "第 6 節導師暫列 Judy / TBC。"
)


def find_unique(base: Path, pattern: str) -> Path:
    matches = [path for path in base.rglob(pattern) if path.is_file()]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one match for {pattern!r} under {base}, found {matches}")
    return matches[0]


def table_rows(path: Path) -> list[list[str]]:
    document = Document(path)
    lesson_table = next(
        table
        for table in document.tables
        if table.rows and table.rows[0].cells[0].text.strip() == "節數"
    )
    return [
        [cell.text.replace("\n", " ").strip() for cell in row.cells]
        for row in lesson_table.rows
    ]


def update_hk239_confirmation(path: Path) -> None:
    document = Document(path)
    paragraph = next(
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.strip().startswith("最新確認：")
    )
    if paragraph.runs:
        paragraph.runs[0].text = HK239_CONFIRMATION
        paragraph.runs[0].bold = True
        paragraph.runs[0].font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        run = paragraph.add_run(HK239_CONFIRMATION)
        run.bold = True
        run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    document.save(path)


def assert_hk239_confirmation(path: Path) -> None:
    document = Document(path)
    matches = [
        paragraph.text.strip()
        for paragraph in document.paragraphs
        if paragraph.text.strip().startswith("最新確認：")
    ]
    assert matches == [HK239_CONFIRMATION], matches


def assert_authoritative_sources(hk239: Path, hk244: Path) -> None:
    hk239_rows = table_rows(hk239)
    expected_239 = {
        "1": ("8月14日", "1000-1300", "Garett"),
        "2": ("8月14日", "1400-1700", "Garett"),
        "3": ("8月19日", "1000-1300", "Garett"),
        "4": ("8月19日", "1400-1700", "Garett（持續評估／小組討論及專題報告）"),
        "5": ("8月21日", "1000-1300", "Judy"),
        "6": ("8月21日", "1400-1700", "Judy / TBC（期末筆試1530-1630）"),
    }
    for row in hk239_rows[1:]:
        lesson = row[0]
        date, time, note = expected_239[lesson]
        assert row[1] == date, row
        assert row[3] == time, row
        assert row[5] == note, row

    hk244_rows = table_rows(hk244)
    expected_244 = {
        "3": ("8月14日", "1400-1800", "Calvin (sit in)"),
        "7": ("8月31日", "1400-1800", "Chan Shuk Ki"),
        "8": ("9月2日", "1400-1730", "Garett｜持續評估小組匯報"),
        "10": ("9月4日", "1400-1730", "Calvin"),
        "11": ("9月7日", "1400-1730", "Calvin｜持續筆試"),
        "12": ("9月8日", "1400-1800", "Calvin｜期末實務試1430-1730"),
    }
    by_lesson = {row[0]: row for row in hk244_rows[1:]}
    for lesson, (date, time, note) in expected_244.items():
        row = by_lesson[lesson]
        assert row[1] == date, row
        assert row[3] == time, row
        assert row[5] == note, row


def sync_one(source: Path, destination: Path) -> None:
    SUPERSEDED.mkdir(parents=True, exist_ok=True)
    if destination.exists():
        backup = SUPERSEDED / f"{destination.stem}_V20k_WRONG{destination.suffix}"
        if not backup.exists():
            shutil.copy2(destination, backup)
    shutil.copy2(source, destination)


def main() -> None:
    course_root = ROOT / "__Course New"
    hk239_source = find_unique(course_root, "*HK239HG_FS_FINAL FINAL*.docx")
    hk244_source = find_unique(course_root, "*HK244HG_CW8_FINAL FINAL*.docx")
    update_hk239_confirmation(hk239_source)
    assert_authoritative_sources(hk239_source, hk244_source)
    assert_hk239_confirmation(hk239_source)

    hk239_destination = find_unique(FINAL, "*HK239HG_FS_FINAL FINAL*.docx")
    hk244_destination = find_unique(FINAL, "*HK244HG_CW8_FINAL FINAL*.docx")
    sync_one(hk239_source, hk239_destination)
    sync_one(hk244_source, hk244_destination)

    assert table_rows(hk239_destination) == table_rows(hk239_source)
    assert table_rows(hk244_destination) == table_rows(hk244_source)
    assert_hk239_confirmation(hk239_destination)
    print(hk239_destination)
    print(hk244_destination)
    print(SUPERSEDED)


if __name__ == "__main__":
    main()
