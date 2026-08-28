from __future__ import annotations

import hashlib
import os
import shutil
import sys
from pathlib import Path

from docx import Document


if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8")


SOURCE = Path(
    r"D:\Garett Super Jobs 2026\Calvin\REAL ERB\Check schedule only - back up on jul 16, don't touch it\HK267HG(CW2)_.docx"
)
WORKFLOW_ROOT = Path(
    r"D:\Garett Super Jobs 2026\Calvin\REAL ERB\Check schedule only (Codex)"
)
FILE_NAME = (
    "2026 11 11 (2026 11 11) - "
    "HK267HG_CW2_CHECKED_AVAILABLE (ENG) - CA - 彩雲二邨.docx"
)
ARCHIVE_NAME = (
    "2026 11 11 (2026 11 11) - "
    "HK267HG_CW2_UNCONFIRMED_ORIGINAL (ENG) - CA - 彩雲二邨.docx"
)

ARCHIVE = WORKFLOW_ROOT / "01 Received - Originals" / "HK267HG__CW2" / ARCHIVE_NAME
SEND_TO_CALVIN = WORKFLOW_ROOT / "03 Send to Calvin" / FILE_NAME
UNCONFIRMED = WORKFLOW_ROOT / "Calvin Unconfirmed - Send to Calvin" / FILE_NAME

EXPECTED_SOURCE_SHA256 = "964db9e71369a5c0c74248498f4726ea37f342acdbb6e8ed4b5750c7903f50a5"
EXPECTED_LESSONS = [
    ("1", "11/11", "Wed", "0900-1300", ""),
    ("2", "12/11", "Thu", "0900-1300", ""),
    ("3", "13/11", "Fri", "0900-1300", ""),
    ("4", "16/11", "Mon", "0900-1300", ""),
    ("5", "17/11", "Tue", "0900-1230", ""),
    ("6", "18/11", "Wed", "0900-1230", "Continous Assessment-Practical Test"),
    ("7", "19/11", "Thu", "0900-1230", ""),
    ("8", "20/11", "Fri", "0900-1230", "Final Practical Exam 1000-1200"),
]


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def copy_original_once() -> None:
    ARCHIVE.parent.mkdir(parents=True, exist_ok=True)
    if ARCHIVE.exists():
        if sha256(ARCHIVE) != EXPECTED_SOURCE_SHA256:
            raise RuntimeError(f"Archive exists with unexpected content: {ARCHIVE}")
        return
    shutil.copy2(SOURCE, ARCHIVE)


def validate_source(document: Document) -> None:
    if len(document.tables) != 2:
        raise RuntimeError(f"Expected 2 tables, found {len(document.tables)}")
    instructor_table, lesson_table = document.tables
    if len(instructor_table.rows) != 2 or len(instructor_table.columns) != 2:
        raise RuntimeError("Unexpected instructor table layout")
    if len(lesson_table.rows) != 9 or len(lesson_table.columns) != 5:
        raise RuntimeError("Unexpected lesson table layout")
    actual = [tuple(cell.text.strip() for cell in row.cells) for row in lesson_table.rows[1:]]
    if actual != EXPECTED_LESSONS:
        raise RuntimeError(f"Lesson table differs from expected source: {actual!r}")


def prepare_reply() -> None:
    if sha256(SOURCE) != EXPECTED_SOURCE_SHA256:
        raise RuntimeError("Source file changed; refusing to prepare a reply")

    source_document = Document(SOURCE)
    validate_source(source_document)
    copy_original_once()

    SEND_TO_CALVIN.parent.mkdir(parents=True, exist_ok=True)
    UNCONFIRMED.parent.mkdir(parents=True, exist_ok=True)
    temporary = SEND_TO_CALVIN.with_suffix(".tmp.docx")
    shutil.copy2(SOURCE, temporary)

    document = Document(temporary)
    validate_source(document)
    instructor_table = document.tables[0]
    instructor_table.cell(1, 0).text = "THK11032"
    instructor_table.cell(1, 1).text = "WONG WAI HON"
    document.save(temporary)
    os.replace(temporary, SEND_TO_CALVIN)
    shutil.copy2(SEND_TO_CALVIN, UNCONFIRMED)

    for output in (SEND_TO_CALVIN, UNCONFIRMED):
        checked = Document(output)
        validate_source(checked)
        values = [cell.text.strip() for cell in checked.tables[0].rows[1].cells]
        if values != ["THK11032", "WONG WAI HON"]:
            raise RuntimeError(f"Instructor fields were not saved correctly: {output}")
        if "internal" in "\n".join(p.text for p in checked.paragraphs).lower():
            raise RuntimeError(f"Internal wording found in Calvin-facing output: {output}")

    if sha256(SOURCE) != EXPECTED_SOURCE_SHA256:
        raise RuntimeError("Source file was modified")
    if sha256(ARCHIVE) != EXPECTED_SOURCE_SHA256:
        raise RuntimeError("Archived original differs from source")

    print(f"source_sha256={sha256(SOURCE).upper()}")
    print(f"archive={ARCHIVE}")
    print(f"reply={SEND_TO_CALVIN}")
    print(f"unconfirmed_copy={UNCONFIRMED}")


if __name__ == "__main__":
    prepare_reply()
